"""Claude API layer: tailored question generation and STAR-method answer scoring."""

import base64
from typing import Literal

import anthropic
from pydantic import BaseModel, Field

MODEL = "claude-opus-4-8"


def make_client(api_key: str) -> anthropic.Anthropic:
    """Build a client for a specific user's key (no shared global client)."""
    return anthropic.Anthropic(api_key=api_key)


def validate_key(api_key: str) -> None:
    """Check a user-entered key is usable.

    Raises ValueError for an empty key, anthropic.AuthenticationError if the key
    is rejected, or anthropic.APITimeoutError / APIConnectionError if Anthropic
    can't be reached. The models endpoint is free and returns 401 on a bad key -
    a cheap probe. Short timeout + single retry so it fails fast, never hangs.
    """
    api_key = (api_key or "").strip()
    if not api_key:
        raise ValueError("Please enter a key.")
    anthropic.Anthropic(api_key=api_key, timeout=15.0, max_retries=1).models.list(limit=1)


# ---------------------------------------------------------------------------
# Structured output models
# ---------------------------------------------------------------------------

class Question(BaseModel):
    id: int
    competency: str = Field(description="The competency this question tests, e.g. 'Stakeholder management'")
    question: str = Field(description="The full competency-based question, phrased as an interviewer would ask it")
    why_asked: str = Field(description="One sentence linking this question to the job spec and/or the candidate's CV")


class QuestionSet(BaseModel):
    questions: list[Question]


class StarComponent(BaseModel):
    present: bool = Field(description="Whether this STAR element was identifiably covered in the answer")
    score: int = Field(description="0-5. 0 = absent, 3 = mentioned but thin, 5 = specific, detailed and convincing")
    feedback: str = Field(description="One or two sentences on how well this element was handled and how to improve it")


class Segment(BaseModel):
    text: str = Field(description="A verbatim chunk of the candidate's answer")
    category: Literal["strong", "expand", "unclear"] = Field(
        description="strong = answered well; expand = right idea but should be developed further; unclear = vague, rambling or hard to follow"
    )
    comment: str = Field(description="Short note on why this chunk got its category, empty string if self-evident")


class Feedback(BaseModel):
    situation: StarComponent
    task: StarComponent
    action: StarComponent
    result: StarComponent
    depth_score: int = Field(description="0-10. How much specific, concrete detail the answer contained overall")
    overall_score: int = Field(description="0-100. Overall quality of the answer for this question at interview")
    segments: list[Segment] = Field(
        description="The candidate's full answer split into consecutive verbatim chunks, in original order, each categorised"
    )
    missed_points: list[str] = Field(
        description="Specific things a strong answer would have included but this one didn't, "
        "including (if a CV was provided) relevant experience visible on the CV that the "
        "candidate failed to bring up"
    )
    suggestions: list[str] = Field(description="Concrete, actionable ways to improve this answer next time")
    summary: str = Field(description="Two or three encouraging but honest sentences summarising the answer")


# ---------------------------------------------------------------------------
# CV handling
# ---------------------------------------------------------------------------

def _docx_to_text(file_bytes: bytes) -> str:
    """Extract readable text from a .docx CV, including tables."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(p for p in parts if p.strip())


def build_cv_block(filename: str, file_bytes: bytes) -> dict:
    """Turn an uploaded CV into a document content block for the API."""
    name = filename.lower()
    if name.endswith(".pdf"):
        source = {
            "type": "base64",
            "media_type": "application/pdf",
            "data": base64.standard_b64encode(file_bytes).decode("ascii"),
        }
    elif name.endswith(".docx"):
        source = {
            "type": "text",
            "media_type": "text/plain",
            "data": _docx_to_text(file_bytes),
        }
    else:  # plain text
        source = {
            "type": "text",
            "media_type": "text/plain",
            "data": file_bytes.decode("utf-8", errors="replace"),
        }
    return {
        "type": "document",
        "source": source,
        "title": "Candidate CV",
        # The CV is resent on every scoring call; cache it so repeat calls are cheap.
        "cache_control": {"type": "ephemeral"},
    }


# ---------------------------------------------------------------------------
# Question generation
# ---------------------------------------------------------------------------

QUESTION_SYSTEM = """You are an experienced interview coach and hiring manager. You design \
competency-based ("tell me about a time...") interview questions.

Given the job spec for a role — and the candidate's CV, when one is provided:
- Identify the key competencies the role requires (e.g. leadership, stakeholder management, \
delivering under pressure, problem solving, communication, adaptability).
- Write one question per competency, each covering a DIFFERENT competency.
- If a CV is provided, tailor questions to the candidate's actual experience where possible, \
so the question gives them a fair chance to draw on something on their CV — but do not answer \
for them. If no CV is provided, base the questions on the job spec alone.
- Phrase each question exactly as a real interviewer would ask it.
- Number ids sequentially from 1."""


def _user_content(cv_block: dict | None, text: str) -> list[dict]:
    """Assemble the user content blocks, with the CV document first when present."""
    blocks: list[dict] = [cv_block] if cv_block else []
    blocks.append({"type": "text", "text": text})
    return blocks


def generate_questions(api_key: str, cv_block: dict | None, job_spec: str, num_questions: int) -> QuestionSet:
    client = make_client(api_key)
    text = (
        f"JOB SPEC:\n{job_spec}\n\n"
        + ("" if cv_block else "No CV was provided - work from the job spec alone.\n\n")
        + f"Generate exactly {num_questions} competency-based interview questions "
        "for this candidate and role."
    )
    response = client.messages.parse(
        model=MODEL,
        max_tokens=16000,
        thinking={"type": "adaptive"},
        system=QUESTION_SYSTEM,
        messages=[{"role": "user", "content": _user_content(cv_block, text)}],
        output_format=QuestionSet,
    )
    return response.parsed_output


# ---------------------------------------------------------------------------
# Answer scoring
# ---------------------------------------------------------------------------

SCORING_SYSTEM = """You are an experienced interview coach assessing a candidate's spoken answer \
to a competency-based interview question, using the STAR method (Situation, Task, Action, Result).

Important context: the answer is an automatic transcript of a SPOKEN answer. Ignore punctuation, \
capitalisation and minor transcription artefacts entirely — judge the substance only. Mild filler \
("um", repeated words) should not be penalised unless it genuinely obscures the point.

Assess:
1. STAR coverage — score each element 0-5. Strong answers set the scene briefly (Situation), make \
their personal responsibility clear (Task), spend most time on specific things THEY did (Action), \
and finish with a concrete, ideally measurable outcome plus reflection (Result).
2. Depth — does the answer contain specific, verifiable detail (names of systems, sizes of teams, \
numbers, timeframes), or could anyone have said it?
3. Relevance — does it actually answer the question and demonstrate the competency being tested?

For `segments`: split the ENTIRE answer into consecutive verbatim chunks in their original order — \
when concatenated they must reproduce the whole transcript. Categorise each chunk: \
"strong" (lands well), "expand" (right idea, needs more development/specificity), \
"unclear" (vague, off-track or hard to follow). Split at natural shifts in quality or topic; \
aim for chunks of one to three sentences.

For `missed_points`: be specific. If a CV is provided, include relevant experience that is \
visible on the CV but the candidate failed to mention — naming what's on the CV — as well as \
generic gaps (no metrics, no reflection, didn't state their personal role, etc.). If no CV is \
provided, focus on the generic gaps and what the question was really probing for.

Be honest but constructive: the goal is that the candidate scores higher next time."""


def score_answer(api_key: str, cv_block: dict | None, job_spec: str, question: dict,
                 transcript: str, clarification: str = "") -> Feedback:
    client = make_client(api_key)
    text = (
        f"JOB SPEC:\n{job_spec}\n\n"
        + ("" if cv_block else "No CV was provided.\n\n")
        + f"COMPETENCY BEING TESTED: {question['competency']}\n"
        f"QUESTION ASKED: {question['question']}\n\n"
        f"CANDIDATE'S SPOKEN ANSWER (auto-transcribed):\n{transcript}"
    )
    if clarification.strip():
        text += (
            "\n\nCANDIDATE'S CLARIFICATION OF THE TRANSCRIPT:\n"
            "The transcript above came from speech-to-text and mis-heard some things. The "
            "candidate has clarified the points below — typically correcting domain terms or "
            "acronyms that were transcribed wrongly. Treat these as corrections to what they "
            "ACTUALLY said, apply them to the answer, and re-evaluate fairly as if those words "
            "had been transcribed correctly. Do not penalise the original mis-transcription. "
            "Only credit substance the candidate genuinely conveyed (corrected wording counts; "
            "brand-new claims that go beyond what was said should not materially inflate the "
            f"score).\n\nClarification:\n{clarification.strip()}"
        )
    response = client.messages.parse(
        model=MODEL,
        max_tokens=16000,
        thinking={"type": "adaptive"},
        system=SCORING_SYSTEM,
        messages=[{"role": "user", "content": _user_content(cv_block, text)}],
        output_format=Feedback,
    )
    return response.parsed_output
